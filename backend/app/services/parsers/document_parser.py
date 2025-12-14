"""
Document Parser Service
Handles Word, Excel, and Markdown files
"""
from docx import Document
import openpyxl
from typing import Dict
import logging

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parser for various document formats"""

    @staticmethod
    def parse_word(file_path: str) -> Dict:
        """
        Parse Word document (.docx)

        Args:
            file_path: Path to Word file

        Returns:
            Dictionary with text and metadata
        """
        try:
            doc = Document(file_path)

            # Extract text from paragraphs
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

            # Extract metadata
            metadata = {
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
                "subject": doc.core_properties.subject or "",
                "keywords": doc.core_properties.keywords or "",
                "created": str(doc.core_properties.created) if doc.core_properties.created else "",
                "modified": str(doc.core_properties.modified) if doc.core_properties.modified else "",
                "paragraphs": len(doc.paragraphs)
            }

            return {"text": text, "metadata": metadata}

        except Exception as e:
            logger.error(f"Error parsing Word document: {e}")
            raise

    @staticmethod
    def parse_excel(file_path: str) -> Dict:
        """
        Parse Excel file (.xlsx)

        Args:
            file_path: Path to Excel file

        Returns:
            Dictionary with text and metadata
        """
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            text = ""

            for sheet in workbook.worksheets:
                text += f"\n\n=== Sheet: {sheet.title} ===\n"

                for row in sheet.iter_rows(values_only=True):
                    # Filter out None values and convert to strings
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    if any(row_values):  # Only add non-empty rows
                        text += "\t".join(row_values) + "\n"

            metadata = {
                "sheets": len(workbook.worksheets),
                "sheet_names": [sheet.title for sheet in workbook.worksheets]
            }

            return {"text": text.strip(), "metadata": metadata}

        except Exception as e:
            logger.error(f"Error parsing Excel file: {e}")
            raise

    @staticmethod
    def parse_markdown(file_path: str) -> Dict:
        """
        Parse Markdown file (.md)

        Args:
            file_path: Path to Markdown file

        Returns:
            Dictionary with text and metadata
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()

            # Extract title from first # heading
            title = ""
            for line in text.split('\n'):
                if line.startswith('# '):
                    title = line[2:].strip()
                    break

            metadata = {
                "title": title,
                "lines": len(text.split('\n'))
            }

            return {"text": text, "metadata": metadata}

        except Exception as e:
            logger.error(f"Error parsing Markdown file: {e}")
            raise

    @staticmethod
    def parse_text(file_path: str) -> Dict:
        """
        Parse plain text file (.txt)

        Args:
            file_path: Path to text file

        Returns:
            Dictionary with text and metadata
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()

            metadata = {
                "lines": len(text.split('\n')),
                "characters": len(text)
            }

            return {"text": text, "metadata": metadata}

        except Exception as e:
            logger.error(f"Error parsing text file: {e}")
            raise
